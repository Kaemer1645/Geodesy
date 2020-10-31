'''import docx

doc=docx.Document('karta.docx')

def printTables(doc):
    wiersze = 0
    kolumny = 0
    for table in doc.tables:
        print('ryz z maslem')
        for row in table.rows:
            wiersze+=1
            for cell in row.cells:
                kolumny+=1
                for paragraph in cell.paragraphs:
                    print(paragraph.text)
                    if paragraph.text == 'nr studzienki':
                        cellka = table.cell(int(wiersze), int(kolumny))
                        cellka.text='1234567'
                        doc.save('kunegunda.docx')

    print('-----')
    print(wiersze)
    print(kolumny)
    row_count = len(table.rows)
    col_count = len(table.columns)
    print([row_count,col_count])
printTables(doc)'''




import docx

doc=docx.Document('karta.docx')
text_file=open('tekstowy.txt')
def printTables(doc):
    nr=0
    dane_w_liniach=[]
    for lines in text_file.readlines():
        #for iteracja in range(0+dupka,10+dupka):   --- zrobic te iteracje po numerach 0-15,,,, 15-25
        dupka=0
        nr += 1
        if nr <11:
            dane_w_liniach.append(lines)
            wiersze = 0
            kolumny = 0
            for table in doc.tables:
                for cols in table.columns:
                    kolumny += 1
                for row in table.rows:
                    wiersze+=1
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            #tu utworzyc zeby bralo wiersz np 2 i zmienialo co 10 - interwal
                            if paragraph.text == 'nr studzienki':
                                cellka = table.cell(2, 1)
                                cellka.text = dane_w_liniach[0]
                            if paragraph.text == 'Opis położenia studzienki':
                                cellka = table.cell(3, 2)
                                cellka.text = dane_w_liniach[1]
                            # doc.save('studzienka_'+str(nr)+'.docx')
            if nr == 10:
                print('Utworzono studzienka_' + '.docx')
                dupka+=15
            #for u in range(0,nr,10):

    print(nr)
    print('-----')
    print(wiersze)
    print(kolumny)
    print('Koniec pracy programu')
    print(dane_w_liniach)
printTables(doc)

