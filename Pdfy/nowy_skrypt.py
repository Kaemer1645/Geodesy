import docx
doc=docx.Document('karta.docx')
txt_file=open('tekstowy.txt')
def printTables(doc):
    linie=[]
    for lines in txt_file.readlines():
        linie.append(lines)
    step=0
    obiekty=[]
    for iteracja in range(0,len(linie),10):
        obiekty.append(linie[step:step+10])
        step+=15
    print(obiekty)
    obiekty2 = [x for x in obiekty if x != []]   #usuwanie pustych list z listy
    print(obiekty2)
    wiersze = 0
    kolumny = 0
    nr=0
    for pojedynczy in obiekty2:
        nr+=1
        for table in doc.tables:
            for cols in table.columns:
                kolumny += 1
            for row in table.rows:
                wiersze += 1
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text == 'nr studzienki':
                            cellka = table.cell(2, 1)
                            cellka.text = pojedynczy[0]
                        if paragraph.text == 'Opis położenia studzienki':
                            cellka = table.cell(3, 2)
                            cellka.text = pojedynczy[1]
        print('Utworzono studzienka_' + str(nr) + '.docx')
        doc.save('studzienka_' + str(nr) + '.docx')
docki=printTables(doc)
